from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "TheBellyDanceDressingRoom-Security-Checklist.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=8.5):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_width(cell, width):
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "DADCE0")


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.bold = True
    run.font.color.rgb = RGBColor(46, 116, 181) if level < 3 else RGBColor(31, 77, 120)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)


rows = [
    ("Secrets", "Exposed database credentials", "Database passwords or service role keys visible in code or public files.", "Covered: no database password/service role key should be in the app. Supabase anon key is browser-safe only with RLS. Action: keep config.js out of GitHub."),
    ("Secrets", "Public .env files", "Environment files accidentally uploaded to GitHub or Netlify output.", "Covered: .gitignore excludes .env and .env.*."),
    ("Secrets", "Hardcoded API keys", "Private keys pasted directly into frontend code.", "Partly covered: config.js has only Supabase anon/publishable key. Action: never put service role keys, OAuth client secrets, or payment secrets in frontend JavaScript."),
    ("Secrets", "Build logs leaking secrets", "Deployment logs print tokens, passwords, or private keys.", "Covered for now: no build step configured. Action: avoid echoing secrets in future Netlify build commands."),
    ("Secrets", "Leaked GitHub repo or commit history", "Secrets were committed even if later deleted.", "Partly covered: .gitignore added. Action: if a real secret is ever committed, rotate it in Supabase/Google/PayPal."),
    ("Auth", "Weak or missing authentication", "Anyone can perform actions without signing in.", "Covered: Google sign-in through Supabase is required for posting, messaging, reporting, and profile updates."),
    ("Auth", "No authorization checks", "Signed-in users can modify things they do not own.", "Covered: app checks ownership and Supabase RLS requires auth.uid() = owner_id for listing updates/deletes."),
    ("Auth", "Users accessing other users' data", "Private messages, reports, or profiles leak between users.", "Covered for conversations: RLS restricts messages to buyer/seller participants. Public listings are intentionally readable."),
    ("Auth", "Open database read/write permissions", "Anon/authenticated roles can read or write too broadly.", "Improved: supabase-security-hardening.sql narrows several grants. Public listing read remains intentional."),
    ("Auth", "Client-side only security checks", "Protection only hides buttons but backend still allows abuse.", "Covered for listings: owner checks exist in both app code and Supabase RLS."),
    ("Auth", "IDOR", "User changes an ID in a request to access or modify someone else's object.", "Covered for listings: RLS requires owner_id match. Covered for conversations: participant policies check buyer/seller IDs."),
    ("Auth", "APIs trust user-controlled IDs or roles", "Frontend claims admin/owner status and backend believes it.", "Covered: policies use auth.uid() instead of trusting frontend role text."),
    ("Storage", "Misconfigured Supabase/S3 buckets", "Uploads can be read, replaced, or deleted by the wrong users.", "Partly covered: listing photos are public by design; upload/update/delete paths require the user's own folder."),
    ("Storage", "Insecure file uploads", "Users upload scripts, huge files, or disguised files.", "Improved: app rejects non-image files and files over 5 MB before upload. Action: keep Supabase Storage policies enabled."),
    ("Storage", "Path traversal", "User input controls file paths like ../../secret.", "Covered: app generates storage paths from signed-in user ID and listing ID, not raw path input."),
    ("Frontend", "XSS", "User text runs as script in other users' browsers.", "Improved: key dynamic text uses textContent. Action: avoid future innerHTML with user-controlled listing/message text."),
    ("Frontend", "CSRF", "A malicious site tricks a signed-in browser into submitting changes.", "Low risk here: no custom cookie-backed server forms. Supabase browser auth handles requests with client library calls."),
    ("Frontend", "Verbose error messages", "Stack traces or internal details are shown to users.", "Partly covered: no stack traces, but Supabase messages still show during beta. Action: replace with friendlier generic messages before launch."),
    ("Frontend", "Source maps exposed", "Production source maps make code easier to inspect.", "Covered: .gitignore excludes *.map. Action: do not deploy source maps publicly."),
    ("Frontend", "Missing security headers", "Browser lacks protections against framing, sniffing, and broad script sources.", "Improved: _headers adds CSP, frame blocking, referrer policy, permissions policy, and nosniff."),
    ("Frontend", "Cookies missing HttpOnly/Secure/SameSite", "Session cookies are easier to steal or misuse.", "Review needed: Supabase browser auth manages sessions. Action: review Supabase Auth/session settings before public launch."),
    ("Frontend", "Overly permissive CORS", "Untrusted domains can interact with auth flows.", "Action: keep Supabase Auth redirect URLs limited to real Netlify/custom domains."),
    ("Backend", "SQL injection", "User input changes raw SQL behavior.", "Covered: app uses Supabase client methods instead of handwritten SQL strings from user input."),
    ("Backend", "NoSQL injection", "User input changes NoSQL query behavior.", "Not applicable: no NoSQL database."),
    ("Backend", "SSRF", "Server fetches attacker-controlled URLs and exposes internal services.", "Not applicable: no server-side URL fetching."),
    ("Backend", "Admin routes unprotected", "Admin screens or APIs are visible to normal users.", "Not applicable yet: no admin route exists. Action: add server-side role checks/RLS before adding admin tools."),
    ("Backend", "Debug pages exposed", "Test pages reveal internals in production.", "Covered: no debug pages found."),
    ("Backend", "Public test/staging environment", "Beta site is accessible to anyone with the URL.", "Partly covered: noindex discourages search indexing. Action: keep link private or add Netlify password protection."),
    ("Backend", "Default credentials unchanged", "Default passwords/secrets remain active.", "Action: keep Supabase database password and Google OAuth secrets private; use strong passwords and account 2FA."),
    ("Backend", "Webhook endpoints without signatures", "Fake payment/webhook events can be submitted.", "Not applicable yet. Required if PayPal/Stripe webhooks are added."),
    ("Payments", "Payment checks only in frontend", "Users can bypass payment/subscription restrictions.", "Current design: payments are direct buyer/seller arrangements. Action: if in-app payments are added, verify payment server-side."),
    ("Abuse", "Missing rate limits", "Attackers spam login, signup, messages, or feedback.", "Partly covered by Supabase platform protections. Action: monitor beta and consider captcha/invite-only testing."),
    ("Privacy", "Logs contain private data", "Tokens, emails, messages, passwords, or payment handles appear in logs.", "Action: avoid console logging secrets or private messages; keep Netlify/Supabase logs access-limited."),
    ("Dependencies", "Dependency vulnerabilities", "Old packages or CDN code contain known vulnerabilities.", "Low current surface: no local npm dependencies. Action: review/update Supabase JS CDN version before production."),
    ("AI", "Prompt injection", "AI features are tricked into leaking or changing data.", "Not applicable: no AI feature exists."),
    ("AI", "AI tools/actions without permission checks", "AI feature can act on data without user authorization.", "Not applicable. If added later, enforce the same RLS/permission checks."),
    ("Database", "Excessive database permissions", "The app role can do more than it needs.", "Improved: supabase-security-hardening.sql narrows grants for reports, feedback, and inquiry status updates."),
    ("Operations", "No audit logs", "No record of who changed or moderated data.", "Not implemented. Action: add audit/moderation tables before wider launch."),
    ("Operations", "No monitoring or alerting", "Abuse or failures go unnoticed.", "Not implemented. Action: monitor Supabase logs and Netlify deploy logs during beta."),
    ("Operations", "No backup or restore plan", "Data cannot be recovered after mistakes or incidents.", "Not implemented. Action: set a recurring Supabase backup/export plan before launch."),
    ("Admin", "Public internal dashboards", "Supabase, Netlify, or Google Cloud dashboards are publicly reachable.", "Covered by separate account logins. Action: protect those accounts with strong passwords and 2FA."),
    ("Privacy", "Unencrypted sensitive data", "Sensitive data stored or sent without encryption.", "Partly covered: Supabase uses HTTPS/in-transit encryption. Action: do not store sensitive payment credentials; only public payment handles are collected."),
    ("Multi-user", "Poor tenant isolation", "One user's private records blend with another's.", "Covered in key flows: RLS separates listing owners and conversation participants. Action: test with two accounts after every RLS change."),
    ("Process", "Over-trusting generated code", "Generated changes are shipped without review/testing.", "Covered by process: changelog, checklist, syntax checks, and two-account testing. Continue reviewing before public launch."),
]


doc = Document()
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Inches(11)
section.page_height = Inches(8.5)
section.top_margin = Inches(0.55)
section.bottom_margin = Inches(0.55)
section.left_margin = Inches(0.55)
section.right_margin = Inches(0.55)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(9)
styles["Heading 1"].font.name = "Calibri"
styles["Heading 1"].font.size = Pt(16)
styles["Heading 1"].font.color.rgb = RGBColor(46, 116, 181)
styles["Heading 2"].font.name = "Calibri"
styles["Heading 2"].font.size = Pt(13)
styles["Heading 2"].font.color.rgb = RGBColor(46, 116, 181)

title = doc.add_paragraph()
title_run = title.add_run("TheBellyDanceDressingRoom Security Checklist")
title_run.bold = True
title_run.font.name = "Calibri"
title_run.font.size = Pt(20)
title_run.font.color.rgb = RGBColor(14, 63, 68)
title.paragraph_format.space_after = Pt(3)

subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run("Records copy | Last reviewed: 2026-08-24")
subtitle_run.font.name = "Calibri"
subtitle_run.font.size = Pt(10)
subtitle_run.font.color.rgb = RGBColor(85, 85, 85)
subtitle.paragraph_format.space_after = Pt(8)

summary = doc.add_paragraph()
summary.add_run("Scope: ").bold = True
summary.add_run(
    "Static Netlify frontend with Supabase Auth, Supabase Database, and Supabase Storage. "
    "The Supabase anon/publishable key is public by design, so the important protection is Row Level Security, narrow grants, storage policies, and careful deployment habits."
)

add_heading(doc, "Quick Actions Before Wider Release", 1)
for item in [
    "Run pending Supabase migrations, especially supabase-security-hardening.sql and supabase-listing-ownership.sql.",
    "Keep config.js out of GitHub; use config.example.js as the shareable template.",
    "Test with two separate Google accounts after every RLS or messaging change.",
    "Keep the beta link private or add Netlify password protection until moderation is stronger.",
    "Create a Supabase backup/restore plan and review Auth redirect URL settings.",
]:
    add_bullet(doc, item)

add_heading(doc, "Checklist Table", 1)
table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"
set_table_borders(table)

headers = ["Area", "Checklist item", "What it means", "How this app covers it / next action"]
widths = [Inches(1.05), Inches(2.05), Inches(2.75), Inches(4.05)]
for idx, header in enumerate(headers):
    cell = table.rows[0].cells[idx]
    set_cell_width(cell, widths[idx])
    set_cell_shading(cell, "E8EEF5")
    set_cell_text(cell, header, bold=True, color="0B2545", size=8.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

for area, item, meaning, coverage in rows:
    cells = table.add_row().cells
    values = [area, item, meaning, coverage]
    for idx, value in enumerate(values):
        set_cell_width(cells[idx], widths[idx])
        set_cell_text(cells[idx], value, size=8)
        cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

doc.add_page_break()
add_heading(doc, "Files Referenced", 1)
for item in [
    ".gitignore",
    "_headers",
    "app.js",
    "SECURITY-CHECKLIST.md",
    "CHANGELOG.md",
    "supabase-security-hardening.sql",
    "supabase-listing-ownership.sql",
    "supabase-schema.sql",
    "supabase-feedback.sql",
    "supabase-filters-messages.sql",
    "supabase-inbox-status.sql",
    "supabase-conversations.sql",
]:
    add_bullet(doc, item)

add_heading(doc, "Notes", 1)
for item in [
    "This document is a practical app security checklist, not a formal penetration test or legal compliance certification.",
    "Public listing photos are intentional for the marketplace experience; private user data should not be placed in public listing text.",
    "If in-app payments, admin dashboards, or AI features are added later, repeat this review because the risk profile changes.",
]:
    add_bullet(doc, item)

doc.save(OUTPUT)
print(OUTPUT)
