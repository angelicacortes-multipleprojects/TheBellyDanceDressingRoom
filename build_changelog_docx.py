from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor


SOURCE = Path("CHANGELOG.md")
OUTPUT = Path("TheBellyDanceDressingRoom-Changelog.docx")


def style_run(run, *, size=10, bold=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


doc = Document()

section = doc.sections[0]
section.top_margin = Pt(54)
section.bottom_margin = Pt(54)
section.left_margin = Pt(54)
section.right_margin = Pt(54)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(10)
styles["Heading 1"].font.name = "Calibri"
styles["Heading 1"].font.size = Pt(18)
styles["Heading 1"].font.color.rgb = RGBColor(14, 63, 68)
styles["Heading 2"].font.name = "Calibri"
styles["Heading 2"].font.size = Pt(14)
styles["Heading 2"].font.color.rgb = RGBColor(46, 116, 181)
styles["Heading 3"].font.name = "Calibri"
styles["Heading 3"].font.size = Pt(11)
styles["Heading 3"].font.color.rgb = RGBColor(31, 77, 120)

title = doc.add_paragraph()
title_run = title.add_run("TheBellyDanceDressingRoom Changelog")
style_run(title_run, size=20, bold=True, color="0E3F44")
title.paragraph_format.space_after = Pt(2)

subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run("Word records copy generated from CHANGELOG.md")
style_run(subtitle_run, size=10, color="66716D")
subtitle.paragraph_format.space_after = Pt(12)

for raw_line in SOURCE.read_text(encoding="utf-8").splitlines():
    line = raw_line.strip()
    if not line:
        continue

    if line.startswith("# "):
        continue
    if line.startswith("## "):
        paragraph = doc.add_paragraph(style="Heading 1")
        paragraph.add_run(line.removeprefix("## ").strip())
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(5)
        continue
    if line.startswith("### "):
        paragraph = doc.add_paragraph(style="Heading 2")
        paragraph.add_run(line.removeprefix("### ").strip())
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(3)
        continue
    if line.startswith("- "):
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(line.removeprefix("- ").strip().replace("`", ""))
        style_run(run, size=9.5)
        continue

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(line.replace("`", ""))
    style_run(run, size=10)

doc.save(OUTPUT)
print(OUTPUT)
